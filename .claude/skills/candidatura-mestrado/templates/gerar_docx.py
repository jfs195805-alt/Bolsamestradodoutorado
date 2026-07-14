# -*- coding: utf-8 -*-
"""Converte projeto_pesquisa.md em DOCX no formato do edital PPGE 2027:
Times New Roman 12, espacamento 1,5, margens sup/inf 2,5cm e laterais 3cm."""
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "projeto_pesquisa.md"
OUT = "Projeto_Pesquisa_Mestrado_PPGE2027_Maria_Souza_v2.docx"

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(3)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
# ABNT: sem espaco extra entre paragrafos do corpo (o recuo marca o paragrafo)
style.paragraph_format.space_after = Pt(0)

def add_runs(p, text):
    # divide em segmentos **negrito** / normal
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            p.add_run(part)
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)

lines = open(SRC, encoding="utf-8").read().splitlines()
i = 0
in_cover = False
in_refs = False
while i < len(lines):
    line = lines[i].rstrip()
    if not line.strip():
        i += 1
        continue
    if line.startswith("# PÁGINA DE ROSTO"):
        in_cover = True
        i += 1
        continue
    if line.startswith("# ") and in_cover:
        # fim da capa -> quebra de pagina
        doc.add_page_break()
        in_cover = False
        # nao pula: processa este heading abaixo
    if line.startswith("|"):
        # tabela: coleta linhas
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            if not all(re.fullmatch(r"-{3,}", c) for c in cells):
                rows.append(cells)
            i += 1
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                c = t.cell(ri, ci)
                c.text = ""
                p = c.paragraphs[0]
                run = p.add_run(cell)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
                if ri == 0:
                    run.bold = True
                if ci > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(0)
        continue
    if line.startswith("## "):
        p = doc.add_paragraph()
        add_runs(p, line[3:])
        for r in p.runs:
            r.bold = True
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif line.startswith("# "):
        in_refs = line.startswith("# 7. REFER")
        p = doc.add_paragraph()
        add_runs(p, line[2:])
        for r in p.runs:
            r.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif in_cover:
        p = doc.add_paragraph()
        add_runs(p, line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
    elif in_refs:
        # ABNT NBR 6023: espaco simples dentro da referencia, 1 linha entre elas,
        # alinhada a esquerda, sem recuo de primeira linha
        p = doc.add_paragraph()
        add_runs(p, line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)
    else:
        p = doc.add_paragraph()
        add_runs(p, line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not line.startswith(("a)", "b)", "c)", "d)", "e)", "1.", "2.", "3.", "4.", "5.", "6.")):
            p.paragraph_format.first_line_indent = Cm(1.25)
    i += 1

doc.save(OUT)
print("gerado:", OUT)
